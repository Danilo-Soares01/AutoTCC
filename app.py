# app.py
# Autor: Assistente Especialista (Engenheiro Full-Stack, Especialista em IA, Normatizador ABNT/Etec, UX/UI Designer)
# Descrição: Aplicação Web com Streamlit para formatar TCCs automaticamente usando a API Gemini.

import streamlit as st
import google.generativeai as genai
import pdfplumber
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import io
import json
import re

# ======================================================================================
# 1. CONFIGURAÇÃO DA PÁGINA E ESTILOS (UX/UI DESIGNER)
# ======================================================================================

st.set_page_config(
    page_title="AutoTCC Formatter",
    page_icon="✨",
    layout="centered",
    initial_sidebar_state="expanded"
)

# Injeção de CSS customizado para uma UI mais suave e moderna
st.markdown("""
<style>
    .main { background-color: #f5f5f5; }
    .stButton>button {
        border-radius: 20px;
        border: 1px solid #007bff;
        background-color: #007bff;
        color: white;
        padding: 10px 24px;
        font-weight: bold;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: white;
        color: #007bff;
    }
    .stFileUploader {
        border: 2px dashed #cccccc;
        border-radius: 10px;
        padding: 20px;
        background-color: #ffffff;
    }
</style>
""", unsafe_allow_html=True)


# ======================================================================================
# 2. FUNÇÕES DE BACK-END (ENGENHEIRO DE SOFTWARE)
# ======================================================================================

def extract_text_from_pdf(file):
    try:
        with pdfplumber.open(file) as pdf:
            full_text = "".join(page.extract_text() for page in pdf.pages if page.extract_text())
        return full_text
    except Exception as e:
        st.error(f"Erro ao ler o arquivo PDF: {e}")
        return None

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return full_text
    except Exception as e:
        st.error(f"Erro ao ler o arquivo DOCX: {e}")
        return None

def extract_text_from_txt(file):
    try:
        return file.getvalue().decode("utf-8")
    except Exception as e:
        st.error(f"Erro ao ler o arquivo TXT: {e}")
        return None

def chunk_text(text, max_chars=15000):
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 < max_chars:
            current_chunk += para + "\n\n"
        else:
            chunks.append(current_chunk)
            current_chunk = para + "\n\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def get_structured_text_from_gemini(api_key: str, raw_text: str):
    try:
        # 1. Configura a chave de API
        genai.configure(api_key=api_key)
        
        # --- ETAPA DE AUTODESCOBERTA ---
        modelos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        modelo_final = next((m for m in modelos if 'flash' in m.lower()), modelos[0])
        model = genai.GenerativeModel(modelo_final)
        # -------------------------------

        # 2. Prompt Completo com as Regras ABNT/Etec (VITAL PARA NÃO DAR ERRO)
        prompt = f"""
        Você é um assistente de formatação acadêmica especialista nas normas ABNT e, principalmente, nas diretrizes institucionais da Etec/Centro Paula Souza. Sua tarefa é analisar o texto bruto fornecido e reestruturá-lo em um formato JSON.

        REGRAS OBRIGATÓRIAS:
        1. O formato de saída deve ser OBRIGATORIAMENTE um JSON válido contendo uma lista de dicionários.
        2. Cada dicionário na lista deve ter duas chaves: "tipo" e "texto".
        3. A chave "tipo" pode ter um dos seguintes valores: "titulo_1", "titulo_2", "paragrafo", "citacao_longa", ou "referencia".
        4. REGRA INSTITUCIONAL ETEC: Converta todas as chamadas de citação para o sistema autor-data. Ex: (SOBRENOME, ANO, p. XX).
        5. Identifique citações com mais de 3 linhas e classifique-as como "citacao_longa", sem aspas.
        6. Identifique a seção de referências bibliográficas. Cada entrada deve ser um item separado com "tipo": "referencia". Ordene alfabeticamente.

        --- INÍCIO DO TEXTO BRUTO ---
        {raw_text}
        --- FIM DO TEXTO BRUTO ---
        """

        # 3. Trava de Segurança: Obriga o Google a responder EXATAMENTE em formato JSON
        configuracao = genai.types.GenerationConfig(
            temperature=0.1, # Temperatura baixa para a IA não "inventar" coisas
            response_mime_type="application/json"
        )
        
        # 4. Chama a geração de conteúdo
        response = model.generate_content(prompt, generation_config=configuracao)
        
        # 5. Limpa a resposta (remove aspas do markdown caso a IA envie) para evitar quebra do JSON
        json_limpo = response.text.strip().replace("```json", "").replace("```", "")
        return json.loads(json_limpo)

    except json.JSONDecodeError:
        st.error("Erro Crítico: A IA não conseguiu estruturar o texto. Tente enviar um trecho menor.")
        return None
    except Exception as e:
        # Mostra exatamente qual modelo tentou usar antes de falhar
        st.error(f"Erro na conexão com a IA ({modelo_final if 'modelo_final' in locals() else 'Desconhecido'}): {e}")
        return None

def validate_references(structured_data):
    cited_authors = set()
    text_body = " ".join([item['texto'] for item in structured_data if item['tipo'] != 'referencia'])
    matches = re.findall(r'\(([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÇ;\s]+),\s*\d{4}', text_body.upper())
    for match in matches:
        authors = [name.strip() for name in match.split(';')]
        for author in authors:
            cited_authors.add(author.split()[-1])

    reference_authors = set()
    references_text = " ".join([item['texto'] for item in structured_data if item['tipo'] == 'referencia'])
    ref_matches = re.findall(r'^([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÇ\s]+),', references_text, re.MULTILINE)
    for match in ref_matches:
        reference_authors.add(match.strip())

    missing_references = cited_authors - reference_authors
    return list(missing_references)

def create_formatted_docx(data_json: list, template_file=None):
    try:
        if template_file:
            document = docx.Document(template_file)
        else:
            document = docx.Document()
            section = document.sections[0]
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.orientation = WD_ORIENT.PORTRAIT
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)

        for item in data_json:
            tipo = item.get("tipo", "paragrafo")
            texto = item.get("texto", "")

            if tipo == "titulo_1":
                p = document.add_paragraph()
                run = p.add_run(texto.upper())
                run.font.name = 'Arial'; run.font.size = Pt(12); run.bold = True
                p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
            elif tipo == "titulo_2":
                p = document.add_paragraph()
                run = p.add_run(texto)
                run.font.name = 'Arial'; run.font.size = Pt(12); run.bold = True
                p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
            elif tipo == "citacao_longa":
                p = document.add_paragraph()
                run = p.add_run(texto)
                run.font.name = 'Arial'; run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(4); p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
            elif tipo == "referencia":
                p = document.add_paragraph()
                run = p.add_run(texto)
                run.font.name = 'Arial'; run.font.size = Pt(12)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(12)
            else:
                p = document.add_paragraph()
                run = p.add_run(texto)
                run.font.name = 'Arial'; run.font.size = Pt(12)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.space_after = Pt(6)

        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        return doc_io

    except Exception as e:
        st.error(f"Erro ao criar o documento DOCX: {e}")
        return None

# ======================================================================================
# 3. INTERFACE PRINCIPAL DA APLICAÇÃO (UI/UX DESIGNER)
# ======================================================================================

with st.sidebar:
    st.image("https://www.etecitanhaem.com.br/wp-content/uploads/2022/02/centro-paula-souza-logo.png", width=150)
    st.header("Configuração")
    
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.success("API Key carregada dos segredos.", icon="✅")
    except (KeyError, FileNotFoundError):
        st.info("Sua chave de API não é armazenada. É usada apenas para esta sessão.")
        api_key = st.text_input("Insira sua Chave de API do Google Gemini", type="password")
    
    st.divider()
    st.header("Template (Opcional)")
    template_file = st.file_uploader("Carregue um template .docx (capa, folha de rosto, etc.)", type=["docx"])
    st.divider()
    st.subheader("Sobre")
    st.markdown("Esta ferramenta foi desenvolvida para auxiliar estudantes da Etec e do Centro Paula Souza na formatação de seus trabalhos acadêmicos.")

st.title("✨ AutoTCC - Formatador ABNT e Etec")
st.subheader("Envie seu trabalho e deixe a IA cuidar da formatação para você.")
st.divider()

tab1, tab2 = st.tabs(["Ferramenta de Formatação", "Guia Rápido e Regras"])

with tab1:
    st.header("1. Envie seu documento")
    uploaded_file = st.file_uploader(
        "Selecione um arquivo (.pdf, .docx, .txt)",
        type=["pdf", "docx", "txt"],
        label_visibility="collapsed"
    )

    if uploaded_file is not None:
        st.header("2. Inicie a formatação")
        st.info(f"Arquivo **{uploaded_file.name}** carregado. Clique no botão abaixo para começar.")

        if st.button("Formatar Documento", use_container_width=True):
            if not api_key:
                st.error("Por favor, insira sua chave de API do Gemini na barra lateral.")
            else:
                # Limpa a sessão anterior para evitar conflito visual de arquivos velhos
                for key in ['doc_pronto', 'missing', 'data', 'file_name']:
                    if key in st.session_state:
                        del st.session_state[key]

                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text("Extraindo texto do documento...")
                raw_text = ""
                file_type = uploaded_file.type
                if file_type == "application/pdf":
                    raw_text = extract_text_from_pdf(uploaded_file)
                elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    raw_text = extract_text_from_docx(uploaded_file)
                elif file_type == "text/plain":
                    raw_text = extract_text_from_txt(uploaded_file)
                progress_bar.progress(10)

                if raw_text:
                    text_chunks = chunk_text(raw_text)
                    num_chunks = len(text_chunks)
                    structured_data_completa = []
                    
                    for i, chunk in enumerate(text_chunks):
                        status_text.text(f"Processando com IA... (Parte {i+1}/{num_chunks})")
                        chunk_data = get_structured_text_from_gemini(api_key, chunk)
                        if chunk_data:
                            structured_data_completa.extend(chunk_data)
                        progress_bar.progress(10 + int((i + 1) / num_chunks * 70))
                    
                    if structured_data_completa:
                        status_text.text("Validando citações e referências...")
                        missing = validate_references(structured_data_completa)
                        progress_bar.progress(85)

                        status_text.text("Montando o novo documento .docx formatado...")
                        formatted_doc = create_formatted_docx(structured_data_completa, template_file)
                        progress_bar.progress(100)
                        
                        st.session_state['doc_pronto'] = formatted_doc
                        st.session_state['missing'] = missing
                        st.session_state['data'] = structured_data_completa
                        st.session_state['file_name'] = f"formatado_{uploaded_file.name.rsplit('.', 1)[0]}.docx"
                        
                        status_text.empty()
                        progress_bar.empty()

        # Mostra o resultado do processamento se existir na sessão
        if 'doc_pronto' in st.session_state:
            st.success("Seu documento foi formatado com sucesso!")
            
            with st.expander("Ver Resumo da Estrutura Gerada", expanded=True):
                data = st.session_state['data']
                counts = {
                    "Títulos": len([i for i in data if i['tipo'] in ['titulo_1', 'titulo_2']]),
                    "Parágrafos": len([i for i in data if i['tipo'] == 'paragrafo']),
                    "Citações Longas": len([i for i in data if i['tipo'] == 'citacao_longa']),
                    "Referências": len([i for i in data if i['tipo'] == 'referencia'])
                }
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("Títulos", counts["Títulos"])
                col2.metric("Parágrafos", counts["Parágrafos"])
                col3.metric("Citações Longas", counts["Citações Longas"])
                col4.metric("Referências", counts["Referências"])

            if st.session_state['missing']:
                st.warning(f"**Aviso de Coerência:** Os seguintes autores foram citados no texto, mas não encontrados nas referências: **{', '.join(st.session_state['missing'])}**. Verifique a grafia ou adicione a referência completa.")

            # Container para organizar os botões finais
            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                st.download_button(
                    label="📥 Baixar Documento (.docx)",
                    data=st.session_state['doc_pronto'],
                    file_name=st.session_state['file_name'],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col_btn2:
                # NOVO BOTÃO: Limpa a sessão e reseta a interface
                if st.button("🔄 Formatar Novo Documento", use_container_width=True):
                    for key in ['doc_pronto', 'missing', 'data', 'file_name']:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.rerun()

with tab2:
    st.header("Guia de Formatação Aplicada")
    st.markdown("""
    Esta ferramenta aplica automaticamente as seguintes regras, baseadas no **Manual Integrado de Normatização Etec/Centro Paula Souza**:
    - **Página:** A4
    - **Margens:** Superior e Esquerda (3 cm); Inferior e Direita (2 cm). (Ignorado se um template for usado).
    - **Fonte Padrão:** Arial, tamanho 12.
    - **Parágrafos:** Espaçamento 1,5, justificado, com recuo de 1,25 cm na primeira linha.
    - **Títulos de Seção:** Negrito. A IA identifica e classifica os níveis.
    - **Citações Longas (+3 linhas):** Recuo de 4 cm à esquerda, fonte 10, espaçamento simples.
    - **Referências:** Alinhamento à esquerda, espaçamento simples, com espaço entre elas.
    - **Sistema de Chamada:** A IA é instruída a converter as citações para o sistema **autor-data**, obrigatório pela Etec. Ex: `(SILVA, 2022, p. 15)`.
    """)
    st.warning("**Atenção:** A IA é uma ferramenta de auxílio. É fundamental que você revise o documento final para garantir que a estrutura e o conteúdo estejam corretos e de acordo com as expectativas do seu orientador.")